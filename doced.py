from abc import ABC, abstractmethod
import re
import os
import json
from PySide6.QtCore import QUrl
from PySide6.QtWidgets import QWidget, QVBoxLayout
import html
from PySide6.QtWebEngineWidgets import QWebEngineView
import docx
import tempfile
import shutil

URL_PATTERN = re.compile(r"^(https?://|www\.)\S+$", re.IGNORECASE)
CACHE_DIR = "cache"
IMAGE_DIR = os.path.join(CACHE_DIR, "images")

os.makedirs(IMAGE_DIR, exist_ok=True)



# ___________________________DOCUMENT___________________________

class Document:

    def __init__(self):
        self.blocks = {}
        self.inuse = []
        self.redo = []
        self.next_id = 1

        self.filename = None
        self.mode = "replace"    # or "append"

    def add_block(self, block):

        id = self.next_id
        self.next_id += 1
        self.blocks[id] = block
        self.inuse.append(id)
        self.redo.clear()


# _____________________________Cache_________________________________

class Cache:

    def __init__(self):
        self.filename = os.path.join(CACHE_DIR, "cache.json")

    def save(self, document):

        data = {
            "blocks": document.blocks,
            "inuse": document.inuse,
            "next_id": document.next_id,
        }

        with open(self.filename, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4)

    def recover(self, document):
        with open(self.filename, "r", encoding="utf-8") as f:
            data = json.load(f)

        document.blocks = {
            int(k): v
            for k, v in data["blocks"].items()
        }

        document.inuse = data["inuse"]
        document.next_id = data["next_id"]
        document.redo.clear()

    def delete(self):
        shutil.rmtree(CACHE_DIR, ignore_errors=True)
        os.makedirs(IMAGE_DIR, exist_ok=True)

# ____________________________EDITOR_________________________________

class Editor:

    def __init__(self):
        self.document = Document()
        self.cache = Cache()

    def execute(self, command):
        command.execute(self.document)
        self.cache.save(self.document)


    def undo(self):

        if not self.document.inuse:
            return

        id = self.document.inuse.pop()
        self.document.redo.append(id)
        self.cache.save(self.document)


    def redo(self):

        if not self.document.redo:
            return

        id = self.document.redo.pop()
        self.document.inuse.append(id)
        self.cache.save(self.document)


# _____________________________COMMANDS_________________________________

class Command(ABC):

    @abstractmethod
    def execute(self, document):
        pass

class Heading(Command):

    def __init__(self, text, level=1):
        self.text = text
        self.level = level

    def execute(self, document):

        document.add_block({
            "type": "heading",
            "level": self.level,
            "text": self.text
        })

class Paragraph(Command):

    def __init__(self, text):
        self.text = text

    def execute(self, document):

        document.add_block({
            "type":"paragraph",
            "text":self.text
        })


class Image(Command):

    def __init__(self, path):
        self.path = path

    def execute(self, document):

        document.add_block({
            "type": "image",
            "text": self.path
        })


class Bullet(Command):

    def __init__(self, text, symbol="bullet"):
        self.text = text
        self.symbol = symbol

    def execute(self, document):

        document.add_block({
            "type": "bullet",
            "symbol": self.symbol,
            "text": self.text
        })

class Hyperlink(Command):

    def __init__(self, text):
        self.text = text

    def execute(self, document):

            if not URL_PATTERN.match(self.text):
                raise ValueError("Invalid URL format")

            url = self.text if self.text.lower().startswith("http") else "http://" + self.text
            document.add_block({
                "type":"hyperlink",
                "text":url
            })


class PreviewWindow(QWidget):

    def __init__(self):
        super().__init__()

        self.setWindowTitle("Preview")
        self.resize(700, 900)

        self.browser = QWebEngineView()

        layout = QVBoxLayout(self)
        layout.addWidget(self.browser)

        self.temp_file = None

    def show_preview(self, filename):

        self.temp_file = filename

        self.browser.load(
            QUrl.fromLocalFile(filename)
        )

        self.show()

    def closeEvent(self, event):

        if self.temp_file and os.path.exists(self.temp_file):
            os.remove(self.temp_file)

        super().closeEvent(event)



# _____________________________RENDERER________________________________


class WordRenderer:

    def render(self, document, filename, mode="replace"):

        if mode == "append" and os.path.exists(filename):
            doc = docx.Document(filename)
        else:
            doc = docx.Document()

        for block_id in document.inuse:

            block = document.blocks[block_id]

            if block["type"] == "heading":
                self.render_heading(doc, block)

            elif block["type"] == "paragraph":
                self.render_paragraph(doc, block)

            elif block["type"] == "bullet":
                self.render_bullet(doc, block)

            elif block["type"] == "image":
                self.render_image(doc, block)

            elif block["type"] == "hyperlink":
                self.render_hyperlink(doc, block)

        doc.save(filename)

    def render_heading(self, doc, block):
        doc.add_heading(
            block["text"],
            level=block.get("level", 1)
        )

    def render_paragraph(self, doc, block):
        doc.add_paragraph(block["text"])


    def render_bullet(self, doc, block):

        symbol = block.get("symbol", "bullet")

        styles = {
            "bullet": "List Bullet",
            "number": "List Number",
            "circle": "List Bullet",
            "square": "List Bullet",
            "check": "List Bullet",
        }

        style = styles.get(symbol)

        if style:
            doc.add_paragraph(
                block["text"],
                style=style,
            )
        else:
            doc.add_paragraph(block["text"])


    def render_image(self, doc, block):
        if os.path.exists(block["text"]):
            doc.add_picture(block["text"])

    def render_hyperlink(self, doc, block):
        paragraph = doc.add_paragraph()
        paragraph.add_run(block["text"])




class HtmlRenderer:

    def render(self, document):

        body = []
        list_open = False

        def close_list():
            nonlocal list_open
            if list_open:
                body.append("</ul>")
                list_open = False

        for block_id in document.inuse:

            block = document.blocks[block_id]

            if block["type"] == "bullet":

                symbols = {
                    "bullet": "•",
                    "circle": "○",
                    "square": "■",
                    "check": "✓",
                }

                symbol = symbols.get(
                    block.get("symbol", "bullet"),
                    "•",
                )

                body.append(
                    f"<p>{symbol} {html.escape(block['text'])}</p>"
                )

            close_list()

            if block["type"] == "heading":

                level = block.get("level", 1)
                body.append(
                    f"<h{level}>{html.escape(block['text'])}</h{level}>"
                )
            elif block["type"] == "paragraph":

                body.append(
                    f"<p>{html.escape(block['text'])}</p>"
                )

            elif block["type"] == "hyperlink":

                url = html.escape(block["text"])

                body.append(
                    f'<p><a href="{url}">{url}</a></p>'
                )

            elif block["type"] == "image":
                abs_path = os.path.abspath(block["text"])
                img = QUrl.fromLocalFile(abs_path).toString()

                body.append(
                    f'<img src="{img}" style="max-width:100%;"><br>'
                )

        close_list()

        page = f"""
                    <html>
                    <head>
                    <meta charset="utf-8">
                    <style>

                    body {{
                        font-family: Calibri, "Segoe UI", sans-serif;
                        margin: 40px;
                        line-height: 1.5;
                    }}

                    img {{
                        display:block;
                        margin:12px 0;
                    }}

                    </style>
                    </head>

                    <body>

                    {''.join(body)}

                    </body>
                    </html>
                    """

        filename = os.path.join(
            tempfile.gettempdir(),
            "preview.html",
        )

        with open(filename, "w", encoding="utf-8") as f:
            f.write(page)

        return filename


# _________________________USAGE_________________________


# editor.execute(Heading("Physics"))
# editor.execute(Paragraph("Newton"))
# editor.execute(Bullet("Egypt"))
# editor.execute(Image("image.jpg"))
# editor.execute(Hyperlink("https://web3.example.com"))

# editor.undo()
# print(editor.document.blocks)
# print(editor.document.inuse)

# editor.redo()
# print(editor.document.blocks)
# print(editor.document.inuse)
